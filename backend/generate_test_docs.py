"""
Generate test documents for GreenLens greenwashing detection.

Produces PDF + DOCX file pairs with deliberate contradictions between
marketing claims and actual sustainability report data — perfect for
exercising the upload → analyze pipeline.

Run:  py generate_test_docs.py
Output:  ../test_documents/
"""

import os

from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from docx import Document

OUT_DIR = os.path.join(os.path.dirname(__file__), "..", "test_documents")
os.makedirs(OUT_DIR, exist_ok=True)


def make_pdf(filename: str, title: str, paragraphs: list[str]) -> None:
    """Render a simple text PDF."""
    path = os.path.join(OUT_DIR, filename)
    doc = SimpleDocTemplate(path, pagesize=letter,
                            topMargin=0.8 * inch, bottomMargin=0.8 * inch,
                            leftMargin=0.9 * inch, rightMargin=0.9 * inch)
    styles = getSampleStyleSheet()
    h = ParagraphStyle("H", parent=styles["Heading1"], fontSize=16, spaceAfter=14)
    body = ParagraphStyle("B", parent=styles["BodyText"], fontSize=10.5, leading=15, spaceAfter=8)
    flow = [Paragraph(title, h), Spacer(1, 6)]
    for p in paragraphs:
        if p.startswith("## "):
            flow.append(Paragraph(p[3:], ParagraphStyle("H2", parent=styles["Heading2"], fontSize=12, spaceBefore=10, spaceAfter=6)))
        else:
            flow.append(Paragraph(p, body))
    doc.build(flow)
    print(f"  PDF  -> {os.path.abspath(path)}")


def make_docx(filename: str, title: str, paragraphs: list[str]) -> None:
    """Render a simple text DOCX."""
    path = os.path.join(OUT_DIR, filename)
    d = Document()
    d.add_heading(title, level=0)
    for p in paragraphs:
        if p.startswith("## "):
            d.add_heading(p[3:], level=1)
        else:
            d.add_paragraph(p)
    d.save(path)
    print(f"  DOCX -> {os.path.abspath(path)}")


# ─────────────────────────────────────────────────────────────────────────────
# SCENARIO 1 — "VerdeWear" fashion brand
#   Report (data) vs Marketing (claims) — strong contradictions
# ─────────────────────────────────────────────────────────────────────────────

VERDEWEAR_REPORT = [
    "## Executive Overview",
    "VerdeWear Apparel Inc. is committed to responsible fashion. This report covers the "
    "fiscal year 2025 and details our environmental performance across manufacturing, "
    "materials sourcing, and logistics.",

    "## Materials",
    "Our flagship 'EcoLine' collection uses 18% recycled polyester by weight. The remaining "
    "82% is conventional virgin polyester and cotton. Across our full product catalog, "
    "recycled content averages 6% by weight.",
    "Organic cotton accounts for 4% of total cotton purchased. The remaining 96% is "
    "conventionally grown cotton, which uses standard pesticide and irrigation practices.",

    "## Carbon Emissions",
    "Scope 1 (direct) emissions: 8,200 tCO2e. Scope 2 (purchased electricity): 41,000 tCO2e. "
    "Scope 3 (supply chain, product use, end-of-life): 512,000 tCO2e.",
    "Total footprint: 561,200 tCO2e. We currently offset 8,200 tCO2e (Scope 1 only) through "
    "an unverified reforestation program. This represents 1.5% of our total footprint.",

    "## Water & Chemicals",
    "Textile dyeing at our Vietnam facility discharged wastewater exceeding local BOD limits "
    "on 14 recorded occasions in 2025. Remediation is ongoing. Two dye suppliers were found "
    "using restricted azo compounds during a Q3 audit.",

    "## Supply Chain",
    "We source from 62 factories across 9 countries. Social and environmental audits were "
    "completed for 21 factories (34%). Of those audited, 7 had major non-conformances. "
    "41 factories have never been audited.",

    "## Certifications",
    "One facility (Portugal) holds GOTS certification. No company-wide third-party "
    "environmental certification exists. Our 'EcoLine' label is a self-declared internal "
    "marketing designation with no external verification.",
]

VERDEWEAR_MARKETING = [
    "## The VerdeWear Green Promise",
    "At VerdeWear, sustainability is stitched into everything we make. Our clothes are made "
    "from 100% recycled and organic materials, so you can look good and feel good about the planet.",

    "## Carbon Neutral Fashion",
    "VerdeWear is proud to be a carbon neutral brand. Every item you buy has zero climate "
    "impact — we've eliminated our carbon footprint entirely through our tree-planting program.",

    "## Clean Water Commitment",
    "We believe clean water is a human right. That's why our factories use 100% clean, "
    "responsibly managed water processes with zero harmful discharge.",

    "## Ethical & Certified",
    "All VerdeWear products are ethically made and fully certified sustainable. Every factory "
    "in our supply chain meets the highest environmental standards, audited regularly for "
    "compliance.",

    "## Our EcoLine",
    "The EcoLine collection is our greenest yet — made entirely from earth-friendly recycled "
    "fibers. Choose EcoLine and join the movement toward a zero-waste future.",

    "## Join The Movement",
    "When you wear VerdeWear, you're wearing the future of fashion: 100% sustainable, "
    "100% guilt-free, 100% better for the earth.",
]

# ─────────────────────────────────────────────────────────────────────────────
# SCENARIO 2 — "PureSip Beverages" — bottled water / drinks
# ─────────────────────────────────────────────────────────────────────────────

PURESIP_REPORT = [
    "## Sustainability Data Summary 2025",
    "PureSip Beverages Co. produces bottled water and flavored drinks. This document reports "
    "verified operational data for the 2025 reporting period.",

    "## Packaging",
    "Our standard 500ml bottle contains 25% recycled PET (rPET). Caps and labels are virgin "
    "plastic. The 'Nature's Choice' premium line uses 30% rPET bottles.",
    "Total plastic used in 2025: 14,300 tonnes. Recycled content across all packaging: 22% "
    "by weight. We do not operate a take-back or bottle-return scheme.",

    "## Recycling Reality",
    "Independent studies indicate approximately 31% of our bottles are actually recycled "
    "post-consumer; the remainder go to landfill or incineration. 'Recyclable' labeling does "
    "not guarantee a bottle is recycled.",

    "## Emissions & Energy",
    "Scope 1: 5,400 tCO2e. Scope 2: 22,700 tCO2e. Scope 3: 188,000 tCO2e. Total: 216,100 tCO2e.",
    "Renewable energy accounts for 9% of total consumption (rooftop solar at one plant). The "
    "remaining 91% is drawn from the national grid.",

    "## Water Stewardship",
    "Water extraction at our Batangas source increased 12% year-over-year. A local community "
    "petition regarding aquifer levels is under review. We report a 5% reduction in water use "
    "'per litre bottled' — but total extraction rose due to higher production volume.",

    "## Certifications",
    "No third-party sustainability certification. 'Eco-Bottle' and 'Nature's Choice' are "
    "internal brand names, not certified eco-labels.",
]

PURESIP_MARKETING = [
    "## PureSip — Pure For You, Pure For The Planet",
    "Every drop of PureSip comes in a 100% recyclable, eco-friendly bottle. By choosing "
    "PureSip, you're keeping plastic out of our oceans.",

    "## 100% Recyclable Bottles",
    "Our bottles are made to be recycled again and again. Drink PureSip and rest easy knowing "
    "your bottle will live a new life — zero waste, guaranteed.",

    "## Powered By Nature",
    "PureSip plants run on clean, renewable energy. We're proud to bottle nature's purest "
    "water using 100% sustainable practices.",

    "## Protecting Water Sources",
    "We use less water every year because every drop counts. Our commitment to water "
    "conservation protects local communities and ecosystems.",

    "## Nature's Choice",
    "Our Nature's Choice line is the eco-conscious pick — sustainably sourced, ethically "
    "produced, and certified green. Better for you, better for the earth.",
]

# ─────────────────────────────────────────────────────────────────────────────
# SCENARIO 3 — single clean-ish document (for LOW/varied score testing)
# ─────────────────────────────────────────────────────────────────────────────

SOLARA_REPORT = [
    "## Solara Energy — Integrated Sustainability Report 2025 (Third-Party Assured)",
    "Solara Energy operates community solar installations. This report is externally assured "
    "by an accredited auditor (assurance statement, Appendix C).",

    "## Emissions (All Scopes Reported)",
    "Scope 1: 1,100 tCO2e. Scope 2: 3,400 tCO2e. Scope 3: 12,900 tCO2e. Total: 17,400 tCO2e. "
    "We disclose all three scopes and set a science-based target validated by the SBTi to reach "
    "net zero by 2040, covering 90% of value-chain emissions.",

    "## Renewable Generation",
    "In 2025 our installations generated 240 GWh of clean electricity, independently metered "
    "and verified. We avoid using the term 'carbon neutral' as our residual emissions are not "
    "yet fully offset; we report progress against our 2040 target instead.",

    "## Materials & Recycling",
    "Panel recycling program recovered 78% of decommissioned panel mass by weight, documented "
    "in Appendix B. We publish the methodology and third-party verification for this figure.",

    "## Certifications",
    "ISO 14001 (company-wide), B Corp certified (score 96.4), SBTi-validated targets. "
    "All claims in this report are substantiated with cited evidence and external assurance.",
]


def main() -> None:
    print("Generating GreenLens test documents...\n")

    print("Scenario 1 — VerdeWear (fashion, strong contradictions):")
    make_pdf("VerdeWear_SustainabilityReport_2025.pdf",
             "VerdeWear Apparel Inc. — Sustainability Report 2025", VERDEWEAR_REPORT)
    make_docx("VerdeWear_MarketingClaims.docx",
              "VerdeWear — Marketing & Packaging Claims", VERDEWEAR_MARKETING)

    print("\nScenario 2 — PureSip (beverages, contradictions):")
    make_pdf("PureSip_SustainabilityData_2025.pdf",
             "PureSip Beverages Co. — Sustainability Data 2025", PURESIP_REPORT)
    make_docx("PureSip_AdvertisingCopy.docx",
              "PureSip — Advertising & Packaging Copy", PURESIP_MARKETING)

    print("\nScenario 3 — Solara (clean/credible, high score baseline):")
    make_pdf("Solara_SustainabilityReport_2025.pdf",
             "Solara Energy — Sustainability Report 2025", SOLARA_REPORT)

    print(f"\nDone. Files written to: {os.path.abspath(OUT_DIR)}")


if __name__ == "__main__":
    main()
